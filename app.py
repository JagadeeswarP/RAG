import streamlit as st
import faiss
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.docstore.in_memory import InMemoryDocstore
from streamlit.runtime.uploaded_file_manager import UploadedFile
import google.generativeai as genai
from dotenv import load_dotenv
import os


load_dotenv()
GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY")

if not GEMINI_API_KEY:
    st.error(" GOOGLE_API_KEY not found in your .env file.")
else:
    genai.configure(api_key=GEMINI_API_KEY)

def generate_gemini_embedding(text):
    """Generate text embeddings using Gemini."""
    model = "models/embedding-001"
    result = genai.embed_content(model=model, content=text)
    return result["embedding"]


def process_input(input_type, input_data):
    """Reads user input and converts it into FAISS vector store."""
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        elif isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        else:
            raise ValueError("Invalid PDF input")
        text = "".join([page.extract_text() or "" for page in pdf_reader.pages])
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data
        else:
            raise ValueError("Text input must be a string")
    elif input_type == "DOCX":
        if isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        elif isinstance(input_data, BytesIO):
            doc = Document(input_data)
        else:
            raise ValueError("Invalid DOCX input")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        if isinstance(input_data, UploadedFile):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        else:
            raise ValueError("Invalid TXT input")
        documents = text
    else:
        raise ValueError("Unsupported input type")


    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = text_splitter.split_text(documents)


    sample_embedding = np.array(generate_gemini_embedding("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)

    vector_store = FAISS(
        embedding_function=generate_gemini_embedding,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)
    return vector_store


def answer_question(vectorstore, query):
    """Retrieve context from vectorstore and use Gemini to answer."""
    retriever = vectorstore.as_retriever()
    docs = retriever.get_relevant_documents(query)
    context = "\n\n".join([d.page_content for d in docs])

    model = genai.GenerativeModel("gemini-1.5-flash")

    prompt = f"""
You are a helpful AI assistant. Use the following context to answer the question clearly and accurately.

Context:
{context}

Question:
{query}

Answer:
"""
    response = model.generate_content(prompt)
    answer = response.text

    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []
    st.session_state["chat_history"].append((query, answer))

    return answer


def main():
    st.title("🧠 RAG Q&A App using Gemini AI")
    st.caption("Built with LangChain, FAISS, and Google Gemini")

    input_type = st.selectbox("Select Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])
    input_data = None

    if input_type == "Link":
        num_links = st.number_input("Number of Links", min_value=1, max_value=20, step=1)
        input_data = [st.text_input(f"URL {i+1}") for i in range(num_links)]
    elif input_type == "Text":
        input_data = st.text_area("Enter Text")
    elif input_type == "PDF":
        input_data = st.file_uploader("Upload PDF", type=["pdf"])
    elif input_type == "DOCX":
        input_data = st.file_uploader("Upload DOCX", type=["docx", "doc"])
    elif input_type == "TXT":
        input_data = st.file_uploader("Upload TXT", type=["txt"])

    if st.button("Process Input"):
        if input_data:
            with st.spinner("Processing input... ⏳"):
                vectorstore = process_input(input_type, input_data)
                st.session_state["vectorstore"] = vectorstore
                st.session_state["chat_history"] = []
            st.success("✅ Vectorstore created successfully!")
        else:
            st.warning("⚠️ Please provide input before processing.")

    if "vectorstore" in st.session_state:
        query = st.text_input("Ask a question about your data:")
        if st.button("Submit Question") and query:
            try:
                with st.spinner("Generating answer... 💭"):
                    answer = answer_question(st.session_state["vectorstore"], query)
                st.markdown(f"**Answer:** {answer}")
            except Exception as e:
                st.error(f"Error: {e}")

if __name__ == "__main__":
    main()
